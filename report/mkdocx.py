#!/usr/bin/env python
"""mkdocx.py — flatten the spec Markdown into one editable A4 Word document.

    python mkdocx.py

Pipeline:
  1. pandoc maps the Markdown 1:1 into a .docx (headings -> Word styles, tables,
     code blocks, block-quotes, lists all preserved).
  2. python-docx then sets A4, margins, a smaller monospace size for the ASCII
     diagrams, and "page break before" on the section headings so each section
     starts on its own page (matches the per-page layout instinct).

Output: report/spec.docx  -> open in Word, rearrange by hand; or drag into Google
Drive to get a native Google Doc. Re-run this any time to regenerate the content;
do your manual page arranging on a *copy* so a re-run doesn't overwrite it.
"""
import subprocess
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")  # cp874 console can't encode our log glyphs
except Exception:
    pass

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_SECTION

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
SOURCES = [ROOT / "draft5.1-1.md", ROOT / "draft5.1-2.md"]
OUT = HERE / "spec.docx"


def run_pandoc() -> None:
    cmd = [
        "pandoc", *[str(s) for s in SOURCES],
        "-f", "gfm",
        "-t", "docx",
        "--toc", "--toc-depth=2",
        "--wrap=preserve",
        "-o", str(OUT),
    ]
    print(">", " ".join(cmd))
    subprocess.run(cmd, check=True)


def restyle() -> None:
    doc = Document(str(OUT))

    # A4 + comfortable report margins on every section.
    for sec in doc.sections:
        sec.page_width, sec.page_height = Mm(210), Mm(297)
        sec.top_margin, sec.bottom_margin = Mm(20), Mm(18)
        sec.left_margin, sec.right_margin = Mm(20), Mm(20)

    styles = doc.styles

    # Shrink the monospace / code styles so wide ASCII diagrams fit the A4 column.
    for name in ("Source Code", "Verbatim Char", "Code", "macro"):
        if name in [s.name for s in styles]:
            try:
                styles[name].font.size = Pt(8)
                styles[name].font.name = "Consolas"
            except Exception:
                pass

    # Each numbered section starts on a fresh page.
    for name in ("Heading 1", "Heading 2"):
        if name in [s.name for s in styles]:
            styles[name].paragraph_format.page_break_before = True

    doc.save(str(OUT))

    n_tables = len(doc.tables)
    n_paras = len(doc.paragraphs)
    n_h1 = sum(1 for p in doc.paragraphs if p.style.name == "Heading 1")
    n_h2 = sum(1 for p in doc.paragraphs if p.style.name == "Heading 2")
    print(f"  A4 set · {n_paras} paragraphs · {n_tables} tables · "
          f"{n_h1} part-headings · {n_h2} section-headings")


if __name__ == "__main__":
    if not all(s.exists() for s in SOURCES):
        sys.exit(f"missing source(s): {[str(s) for s in SOURCES if not s.exists()]}")
    run_pandoc()
    restyle()
    print(f"wrote {OUT}")
